from flask import Flask, render_template, request, redirect, url_for, session, jsonify
import os
import uuid
import json
import pickle
from werkzeug.utils import secure_filename
import google.generativeai as genai
import PyPDF2
import docx
import cv2
import pytesseract
from PIL import Image
import re
import numpy as np
import datetime
import sys

app = Flask(__name__)
app.secret_key = 'your_secret_key'  # Change in production
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max upload size

# Create a directory to store session data
SESSION_DATA_DIR = os.path.join(os.getcwd(), 'session_data')
os.makedirs(SESSION_DATA_DIR, exist_ok=True)

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Experience level mappings
EXPERIENCE_LEVELS = {
    "entry": "Entry-level (0-1yrs)",
    "beginner": "Beginner (1-3yrs)",
    "intermediate": "Intermediate (3-7yrs)",
    "professional": "Professional (7+yrs)"
}

# File extensions allowed
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'jpg', 'jpeg', 'png'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        text = f"Failed to extract text from PDF: {e}"
    
    # If no text was extracted, provide a fallback
    if not text.strip():
        text = "Unable to extract meaningful text from this PDF. It may be image-based or have security restrictions."
    
    return text

def extract_text_from_docx(docx_path):
    text = ""
    try:
        doc = docx.Document(docx_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        text = f"Failed to extract text from DOCX: {e}"
    
    # If no text was extracted, provide a fallback
    if not text.strip():
        text = "Unable to extract meaningful text from this DOCX file."
    
    return text

def extract_text_from_image(image_path):
    text = ""
    try:
        # Check if tesseract is installed
        pytesseract.get_tesseract_version()
        
        # Process the image
        image = cv2.imread(image_path)
        if image is None:
            return "Failed to read image file. The file may be corrupted."
        
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        
        # Apply some image preprocessing to improve OCR results
        gray = cv2.GaussianBlur(gray, (3, 3), 0)
        gray = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
        
        text = pytesseract.image_to_string(gray)
    except Exception as e:
        print(f"Error extracting text from image: {e}")
        if "TesseractNotFoundError" in str(e):
            text = "Tesseract OCR is not installed or not in PATH. Please install Tesseract to process image files."
        else:
            text = f"Failed to extract text from image: {e}"
    
    # If no text was extracted, provide a fallback
    if not text.strip():
        text = "Unable to extract meaningful text from this image."
    
    return text

def extract_text_from_file(file_path):
    file_extension = file_path.rsplit('.', 1)[1].lower()
    
    if file_extension == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == 'docx':
        return extract_text_from_docx(file_path)
    elif file_extension in ['jpg', 'jpeg', 'png']:
        return extract_text_from_image(file_path)
    else:
        return "Unsupported file format"

def analyze_resume_for_ats(resume_text, job_description, job_type, experience_level):
    # Configure the Gemini API
    api_key = os.environ.get("GEMINI_API_KEY", "AIzaSyASwGzSeJI_G0shdSWt3ukhgJPpqABBadM")  # Replace with your actual key for testing
    if not api_key or api_key == "AIzaSyASwGzSeJI_G0shdSWt3ukhgJPpqABBadM":
        return {
            "error": "Valid GEMINI_API_KEY not configured. Either set it in environment variables or in the code.",
            "ats_score": 75,
            "explanation": "This is a placeholder score since the Gemini API is not configured. In a real analysis, this would be based on comparing your resume to the job description.",
            "key_skills_matched": ["Python", "Flask", "Web Development"],
            "missing_skills": ["API Integration", "Data Analysis"],
            "recommendations": ["Add more specific achievements", "Tailor your resume to the job description"]
        }
    
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-pro')
    
    prompt = f"""
    As an ATS analyzer, evaluate the resume against the job description.
    
    Job Type: {job_type}
    Experience Level: {experience_level}
    
    Job Description:
    {job_description}
    
    Resume:
    {resume_text}
    
    Provide an ATS score from 0-100 with an explanation of the score.
    Extract key skills from both the resume and job description.
    Identify missing skills or qualifications.
    Provide formatting and content recommendations to improve the resume.
    
    Format the response as a JSON object with the following structure:
    {{
        "ats_score": <score>,
        "explanation": "<detailed explanation>",
        "key_skills_matched": ["skill1", "skill2", ...],
        "missing_skills": ["skill1", "skill2", ...],
        "recommendations": ["rec1", "rec2", ...]
    }}
    """
    
    try:
        response = model.generate_content(prompt)
        result = response.text
        # Extract JSON from response
        json_match = re.search(r'```json\n(.*?)\n```', result, re.DOTALL)
        if json_match:
            result = json_match.group(1)
        return json.loads(result)
    except Exception as e:
        return {"error": str(e), "ats_score": 0, "explanation": "Failed to analyze resume"}

def generate_interview_questions(resume_text, job_description, job_type, experience_level):
    # Configure the Gemini API
    api_key = os.environ.get("GEMINI_API_KEY", "your_gemini_api_key_here")  # Replace with your actual key for testing
    if not api_key or api_key == "your_gemini_api_key_here":
        # Return placeholder interview questions if API key is not available
        return {
            "questions": [
                {
                    "question": "Tell me about your experience with Flask and web development.",
                    "category": "Technical Skills",
                    "difficulty": "medium",
                    "rationale": "This question helps assess the candidate's web development experience."
                },
                {
                    "question": "How do you approach parsing different file formats like PDFs and images?",
                    "category": "Technical Skills",
                    "difficulty": "hard",
                    "rationale": "This evaluates understanding of document processing techniques."
                },
                {
                    "question": "Describe a challenging project you worked on and how you overcame obstacles.",
                    "category": "Experience",
                    "difficulty": "medium",
                    "rationale": "This reveals problem-solving abilities and project experience."
                },
                {
                    "question": "How do you keep your technical skills up-to-date?",
                    "category": "Professional Development",
                    "difficulty": "easy",
                    "rationale": "This shows commitment to learning and growth."
                },
                {
                    "question": "What experience do you have with API integration?",
                    "category": "Technical Skills",
                    "difficulty": "medium",
                    "rationale": "This assesses API experience which is relevant to the job."
                },
                {
                    "question": "How do you approach debugging complex issues in your code?",
                    "category": "Problem Solving",
                    "difficulty": "hard",
                    "rationale": "This evaluates troubleshooting methodology."
                },
                {
                    "question": "Tell me about your experience working in team environments.",
                    "category": "Soft Skills",
                    "difficulty": "easy",
                    "rationale": "This assesses teamwork and collaboration abilities."
                },
                {
                    "question": "How do you prioritize tasks when working on multiple projects?",
                    "category": "Project Management",
                    "difficulty": "medium",
                    "rationale": "This reveals organizational and time management skills."
                },
                {
                    "question": "Describe your experience with version control systems like Git.",
                    "category": "Technical Skills",
                    "difficulty": "medium",
                    "rationale": "This is important for collaborative development."
                },
                {
                    "question": "What methods do you use to ensure code quality?",
                    "category": "Best Practices",
                    "difficulty": "medium",
                    "rationale": "This shows attention to quality and best practices."
                },
                {
                    "question": "How do you handle receiving critical feedback?",
                    "category": "Soft Skills",
                    "difficulty": "easy",
                    "rationale": "This reveals adaptability and growth mindset."
                },
                {
                    "question": "Explain a complex technical concept in simple terms.",
                    "category": "Communication",
                    "difficulty": "medium",
                    "rationale": "This tests communication and explanation skills."
                },
                {
                    "question": "What motivates you as a developer?",
                    "category": "Motivation",
                    "difficulty": "easy",
                    "rationale": "This helps understand the candidate's drive and passion."
                },
                {
                    "question": "How do you stay focused during repetitive or monotonous tasks?",
                    "category": "Work Habits",
                    "difficulty": "easy",
                    "rationale": "This reveals discipline and work ethic."
                },
                {
                    "question": "Where do you see yourself professionally in five years?",
                    "category": "Career Goals",
                    "difficulty": "medium",
                    "rationale": "This shows long-term vision and ambition."
                }
            ]
        }
    
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-pro')
    
    prompt = f"""
    As an AI interview bot, analyze the resume and job description to generate 15 interview questions.
    
    Job Type: {job_type}
    Experience Level: {experience_level}
    
    Job Description:
    {job_description}
    
    Resume:
    {resume_text}
    
    Focus on skills, experiences, courses, certifications, and other relevant information from the resume.
    Tailor questions to match the job type (technical vs non-technical) and experience level.
    For technical roles, include appropriate technical questions.
    For higher experience levels, include questions about leadership, complex problem-solving, and strategic thinking.
    
    Format the response as a JSON object with the following structure:
    {{
        "questions": [
            {{
                "question": "<question text>",
                "category": "<category>",
                "difficulty": "<easy|medium|hard>",
                "rationale": "<why this question is relevant based on resume/JD>"
            }},
            ...
        ]
    }}
    """
    
    try:
        response = model.generate_content(prompt)
        result = response.text
        # Extract JSON from response
        json_match = re.search(r'```json\n(.*?)\n```', result, re.DOTALL)
        if json_match:
            result = json_match.group(1)
        return json.loads(result)
    except Exception as e:
        return {"error": str(e), "questions": []}

# Helper functions for session data storage
def save_session_data(session_id, data):
    """Save session data to a file"""
    filepath = os.path.join(SESSION_DATA_DIR, f"{session_id}.pkl")
    with open(filepath, 'wb') as f:
        pickle.dump(data, f)
    return filepath

def load_session_data(session_id):
    """Load session data from a file"""
    filepath = os.path.join(SESSION_DATA_DIR, f"{session_id}.pkl")
    if not os.path.exists(filepath):
        return None
    with open(filepath, 'rb') as f:
        return pickle.load(f)

@app.route('/')
def index():
    current_year = datetime.datetime.now().year
    return render_template('index.html', experience_levels=EXPERIENCE_LEVELS, current_year=current_year)

@app.route('/upload', methods=['POST'])
def upload_file():
    print("Upload function called")
    if 'resume' not in request.files:
        print("No resume file in request")
        return redirect(request.url)
    
    file = request.files['resume']
    print(f"File received: {file.filename}")
    
    if file.filename == '':
        print("Empty filename")
        return redirect(request.url)
    
    if file and allowed_file(file.filename):
        # Get form data
        job_type = request.form.get('job_type')
        experience_level = request.form.get('experience_level')
        job_description = request.form.get('job_description')
        
        print(f"Form variables: {job_type}, {experience_level}")
        
        # Create unique filename
        filename = str(uuid.uuid4()) + '_' + secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        print(f"File saved to: {filepath}")
        
        # Extract text from resume
        resume_text = extract_text_from_file(filepath)
        print(f"Extracted text length: {len(resume_text)}")
        
        # Get ATS analysis
        print("About to get ATS analysis")
        try:
            ats_analysis = analyze_resume_for_ats(
                resume_text,
                job_description,
                job_type,
                EXPERIENCE_LEVELS[experience_level]
            )
            print(f"Analysis received: {ats_analysis.get('ats_score')}")
            
            # Generate a unique session ID
            session_id = str(uuid.uuid4())
            
            # Prepare session data
            session_data = {
                'job_type': job_type,
                'experience_level': experience_level,
                'job_description': job_description,
                'resume_text': resume_text,
                'resume_filepath': filepath,
                'ats_analysis': ats_analysis
            }
            
            # Save to temporary file
            save_session_data(session_id, session_data)
            
            # Store only the session ID in the browser cookie
            session['session_id'] = session_id
            print(f"Session ID created and stored: {session_id}")
            
            return redirect(url_for('results'))
        except Exception as e:
            print(f"Error during analysis: {str(e)}")
            return f"Error during analysis: {str(e)}"
    
    return "Invalid file format. Please upload PDF, DOCX, JPG, JPEG, or PNG files."

@app.route('/results')
def results():
    if 'session_id' not in session:
        print("No session_id in session, redirecting to index")
        return redirect(url_for('index'))
    
    # Load session data
    session_data = load_session_data(session['session_id'])
    if not session_data:
        print("No session data found, redirecting to index")
        return redirect(url_for('index'))
    
    current_year = datetime.datetime.now().year
    return render_template(
        'results.html',
        ats_analysis=session_data['ats_analysis'],
        job_type=session_data['job_type'],
        experience_level=EXPERIENCE_LEVELS[session_data['experience_level']],
        current_year=current_year
    )

@app.route('/interview')
def interview():
    if 'session_id' not in session:
        print("No session_id in session, redirecting to index")
        return redirect(url_for('index'))
    
    # Load session data
    session_data = load_session_data(session['session_id'])
    if not session_data:
        print("No session data found, redirecting to index")
        return redirect(url_for('index'))
    
    # Generate interview questions
    interview_data = generate_interview_questions(
        session_data['resume_text'],
        session_data['job_description'],
        session_data['job_type'],
        EXPERIENCE_LEVELS[session_data['experience_level']]
    )
    
    current_year = datetime.datetime.now().year
    return render_template('interview.html', interview_data=interview_data, current_year=current_year)

@app.route('/api/analyze', methods=['POST'])
def api_analyze():
    data = request.json
    
    if not data or 'resume_text' not in data or 'job_description' not in data:
        return jsonify({"error": "Missing required fields"}), 400
    
    job_type = data.get('job_type', 'technical')
    experience_level = data.get('experience_level', 'entry')
    
    ats_analysis = analyze_resume_for_ats(
        data['resume_text'],
        data['job_description'],
        job_type,
        EXPERIENCE_LEVELS.get(experience_level, 'Entry-level (0-1yrs)')
    )
    
    return jsonify(ats_analysis)

@app.errorhandler(Exception)
def handle_exception(e):
    """Handle any unhandled exception"""
    # Log the error
    print(f"Unhandled exception: {str(e)}")
    
    # Get error details
    error_class = e.__class__.__name__
    error_message = str(e)
    
    if isinstance(e, (FileNotFoundError, PermissionError)):
        # File system errors
        message = f"File system error: {error_message}"
    elif "Tesseract" in error_message:
        # Tesseract-related errors
        message = "OCR error: Tesseract is not properly installed or configured."
    elif isinstance(e, ValueError) and "api_key" in error_message.lower():
        # API key related errors
        message = "API error: Invalid or missing Gemini API key."
    else:
        # Generic error message
        message = f"An error occurred: {error_message}"
    
    return render_template(
        'error.html', 
        error_class=error_class,
        error_message=message,
        current_year=datetime.datetime.now().year
    ), 500

@app.route('/debug')
def debug_info():
    """Show debug information to help diagnose issues"""
    session_data = None
    if 'session_id' in session:
        session_data = load_session_data(session['session_id'])
    
    debug_data = {
        "environment": {
            "gemini_api_key_set": bool(os.environ.get("GEMINI_API_KEY")),
            "upload_folder_exists": os.path.exists(app.config['UPLOAD_FOLDER']),
            "session_data_dir_exists": os.path.exists(SESSION_DATA_DIR),
            "python_version": sys.version,
            "libraries": {
                "flask": flask.__version__,
                "pytesseract_installed": "Yes" if 'pytesseract' in sys.modules else "No",
                "pypdf2_installed": "Yes" if 'PyPDF2' in sys.modules else "No",
                "opencv_installed": "Yes" if 'cv2' in sys.modules else "No",
                "docx_installed": "Yes" if 'docx' in sys.modules else "No",
                "google_generativeai_installed": "Yes" if 'google.generativeai' in sys.modules else "No"
            }
        },
        "session": {
            "keys": list(session.keys()) if session else [],
            "session_id": session.get('session_id', "Not set"),
            "session_data_available": session_data is not None
        }
    }

def score_interview_response(question, answer, job_type, experience_level):
    """Score an interview answer using Gemini API"""
    api_key = os.environ.get("GEMINI_API_KEY", "your_gemini_api_key_here")
    if not api_key or api_key == "your_gemini_api_key_here":
        # Return placeholder scoring if API key isn't available
        return {
            "score": 7,  # Score out of 10
            "feedback": "This is placeholder feedback since the Gemini API is not configured. In a real evaluation, detailed feedback would be provided based on your answer.",
            "strengths": ["Clear communication", "Relevant examples"],
            "areas_for_improvement": ["Add more specific details", "Structure your answer using the STAR method"]
        }
    
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-pro')
    
    prompt = f"""
    As an interview coach, evaluate this response to the following interview question.
    
    Job Type: {job_type}
    Experience Level: {experience_level}
    
    Question: {question}
    
    Answer: {answer}
    
    Provide a score from 1-10 with detailed feedback on the strengths and areas for improvement.
    Format the response as a JSON object with the following structure:
    {{
        "score": <score from 1-10>,
        "feedback": "<detailed evaluation of the answer>",
        "strengths": ["strength1", "strength2", ...],
        "areas_for_improvement": ["improvement1", "improvement2", ...]
    }}
    """
    
    try:
        response = model.generate_content(prompt)
        result = response.text
        # Extract JSON from response
        json_match = re.search(r'```json\n(.*?)\n```', result, re.DOTALL)
        if json_match:
            result = json_match.group(1)
        return json.loads(result)
    except Exception as e:
        return {
            "score": 5,
            "feedback": f"Error evaluating response: {str(e)}",
            "strengths": [],
            "areas_for_improvement": ["Try again with a more specific answer"]
        }

def calculate_overall_score(responses):
    """Calculate overall interview score based on individual responses"""
    if not responses:
        return 0
    
    total_score = sum(response.get("score", 0) for response in responses.values())
    return round(total_score / len(responses))

@app.route('/save_response', methods=['POST'])
def save_response():
    """Save a single interview answer"""
    if 'session_id' not in session:
        return jsonify({"error": "Session expired"}), 401
    
    data = request.json
    if not data or 'question_id' not in data or 'answer' not in data or 'question_text' not in data:
        return jsonify({"error": "Missing required fields"}), 400
    
    question_id = data['question_id']
    answer = data['answer']
    question_text = data['question_text']
    
    # Load session data
    session_data = load_session_data(session['session_id'])
    if not session_data:
        return jsonify({"error": "Session data not found"}), 404
    
    # Score the response
    evaluation = score_interview_response(
        question_text, 
        answer, 
        session_data['job_type'], 
        EXPERIENCE_LEVELS[session_data['experience_level']]
    )
    
    # Initialize responses dictionary if it doesn't exist
    if 'interview_responses' not in session_data:
        session_data['interview_responses'] = {}
    
    # Save the response and evaluation
    session_data['interview_responses'][question_id] = {
        "question": question_text,
        "answer": answer,
        "score": evaluation.get("score", 0),
        "feedback": evaluation.get("feedback", ""),
        "strengths": evaluation.get("strengths", []),
        "areas_for_improvement": evaluation.get("areas_for_improvement", [])
    }
    
    # Save updated session data
    save_session_data(session['session_id'], session_data)
    
    return jsonify({
        "success": True, 
        "evaluation": evaluation,
        "overall_score": calculate_overall_score(session_data['interview_responses']),
        "responses": len(session_data['interview_responses'])
    })

@app.route('/interview_results')
def interview_results():
    """Show interview results page"""
    if 'session_id' not in session:
        return redirect(url_for('index'))
    
    # Load session data
    session_data = load_session_data(session['session_id'])
    if not session_data:
        return redirect(url_for('index'))
    
    # Check if there are any responses
    if 'interview_responses' not in session_data or not session_data['interview_responses']:
        return redirect(url_for('interview'))
    
    # Calculate overall score
    overall_score = calculate_overall_score(session_data['interview_responses'])
    
    current_year = datetime.datetime.now().year
    return render_template(
        'interview_results.html',
        responses=session_data['interview_responses'],
        overall_score=overall_score,
        current_year=current_year
    )
    
    # Check if tesseract is installed
    try:
        tesseract_version = pytesseract.get_tesseract_version()
        debug_data["environment"]["tesseract_version"] = str(tesseract_version)
    except Exception as e:
        debug_data["environment"]["tesseract_error"] = str(e)
    
    return render_template('debug.html', debug_data=debug_data, current_year=datetime.datetime.now().year)

if __name__ == '__main__':
    import flask
    app.run(debug=True)